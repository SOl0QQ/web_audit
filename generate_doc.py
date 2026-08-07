#!/usr/bin/env python3
"""
生成 Web Audit 項目架構文檔（Word 格式）
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_document():
    doc = Document()

    # 設置默認字體
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 標題
    title = doc.add_heading('Web Audit 安全審計系統', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副標題
    subtitle = doc.add_paragraph('基於 LangChain + LLM 的自動化 Web 安全審計流水線')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # 目錄
    doc.add_heading('目錄', 1)
    toc_items = [
        '1. 項目概覽',
        '2. 目錄結構',
        '3. 核心架構',
        '4. 數據流',
        '5. 配置系統',
        '6. 核心模塊詳解',
        '7. 審計模塊詳解',
        '8. 報告系統',
        '9. 關鍵技術實現',
        '10. 使用說明',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # 1. 項目概覽
    doc.add_heading('1. 項目概覽', 1)
    doc.add_paragraph(
        'Web Audit 是一個基於 LangChain 和大語言模型（LLM）的自動化 Web 安全審計系統。'
        '它通過智能流水線的方式，自動完成從登錄頁識別、SQL 注入繞過、上傳點發現到 Webshell 驗證的完整攻擊鏈。'
    )

    doc.add_heading('1.1 核心特性', 2)
    features = [
        '多階段攻擊鏈：登錄識別 → SQL 注入繞過 → 上傳點發現 → Webshell 驗證',
        'LLM 驅動：使用大模型進行語義分析，替代傳統規則引擎',
        '多 LLM 支持：支持 Google Gemini、OpenAI GPT、本地 Ollama 模型',
        '動態渲染：集成 Playwright 處理 SPA/Vue/React 等動態頁面',
        '外部工具整合：整合 Katana 爬蟲和 Dirsearch 目錄爆破',
        '智能去重：基於 URL 簽名的智能去重機制',
        '多語言支持：支持中、英、日、韓、西、法、俄、德等多語言',
        '線程安全：支持多目標並發掃描',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('1.2 技術棧', 2)
    tech_stack = [
        'Python 3.10+',
        'LangChain 1.3.2 - LLM 應用框架',
        'Playwright 1.60.0 - 動態頁面渲染',
        'BeautifulSoup4 - HTML 解析',
        'Requests - HTTP 請求',
        'Pydantic - 數據驗證',
        'Katana - 主動爬蟲（可選）',
        'Dirsearch - 目錄爆破（可選）',
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_page_break()

    # 2. 目錄結構
    doc.add_heading('2. 目錄結構', 1)

    structure = """web_audit/
├── main.py                          # 流水線主入口
├── __init__.py                      # 包初始化
│
├── config/                          # 配置模塊
│   ├── __init__.py
│   ├── settings.py                  # 全局配置（LLM、HTTP、爬蟲、工具等）
│   └── login_wordlist.txt           # dirsearch 登錄頁專用字典（~84條）
│
├── core/                            # 核心基礎設施
│   ├── __init__.py
│   ├── requester.py                 # 統一 HTTP 請求封裝（requests + Playwright）
│   ├── parser.py                    # HTML 結構化解析器（BeautifulSoup）
│   ├── logger.py                    # 線程安全日誌緩衝代理
│   ├── llm_factory.py              # LLM 工廠（多 Provider 切換）
│   ├── tool_discovery.py           # 外部工具 URL 發現層（Katana + Dirsearch）
│   ├── playwright_interceptor.py   # Playwright 動態無頭瀏覽器攔截器
│   └── playwright_crawler.py       # Playwright 模擬點擊爬蟲（Phase 4）
│
├── modules/                         # 審計功能模塊
│   ├── __init__.py
│   ├── base_module.py              # 抽像基類（模塊接口契約）
│   ├── login_detector.py           # 模塊 1：登錄頁面識別（三層架構）
│   ├── sqli_detector.py            # 模塊 2：Auth Bypass (SQL 注入繞過)
│   ├── upload_auditor.py           # 模塊 3：文件上傳功能識別
│   └── unified_upload_auditor.py   # 模塊 4：Webshell 上傳繞過與 RCE 驗證
│
└── reports/                         # 報告輸出
    ├── __init__.py
    └── reporter.py                  # 報告生成器"""

    p = doc.add_paragraph()
    p.add_run(structure).font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # 3. 核心架構
    doc.add_heading('3. 核心架構', 1)

    doc.add_heading('3.1 四階段攻擊鏈', 2)
    doc.add_paragraph(
        '系統採用四階段攻擊鏈設計，每個階段獨立可插拔，可單獨運行或組合執行：'
    )

    stages = [
        ('Stage 1: 登錄頁識別', 'LoginDetectorModule',
         '通過外部工具（Katana + Dirsearch）進行廣度發現，然後使用 LLM 進行語義過濾，'
         '識別出真實的登錄頁面。支持三層架構：工具發現 → LLM 過濾 → 遞歸補充。'),
        ('Stage 2: SQL 注入繞過', 'SQLiDetectorModule',
         '對登錄表單進行 Auth Bypass 測試。使用 LLM 生成候選 Payload，'
         '並通過基線響應與 Payload 響應的語義對比，判斷是否繞過成功。'),
        ('Stage 3: 上傳點發現', 'UploadIdentifierModule',
         '在已認證狀態下，通過四階段探索（靜態解析 → 關鍵詞鏈接 → BFS 遍歷 → Katana 動態爬蟲 → Playwright 模擬點擊）'
         '發現所有上傳表單。'),
        ('Stage 4: Webshell 驗證', 'UnifiedUploadAuditModule',
         '使用 LLM Agent 模式進行智能決策：動態生成繞過策略 → 上傳探針 → 路徑推斷 → 診斷與自糾錯循環。'),
    ]

    for stage_name, module_name, desc in stages:
        doc.add_heading(stage_name, 3)
        doc.add_paragraph(f'模塊: {module_name}')
        doc.add_paragraph(desc)

    doc.add_heading('3.2 模塊架構', 2)
    doc.add_paragraph(
        '所有審計模塊都繼承自 BaseModule 抽像基類，實現統一的 run() 接口：'
    )

    code = """class BaseModule(ABC):
    def __init__(self, requester: Requester):
        self.requester = requester

    @abstractmethod
    def run(self, url: str, context: Dict[str, Any]) -> Dict[str, Any]:
        '''
        執行審計邏輯

        Args:
            url: 當前分析的目標 URL
            context: 上下文信息

        Returns:
            標準化的審計結果字典
        '''
        pass"""

    p = doc.add_paragraph()
    p.add_run(code).font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # 4. 數據流
    doc.add_heading('4. 數據流', 1)

    doc.add_heading('4.1 完整攻擊鏈數據流', 2)

    flow = """run_pipeline(target_url)
  │
  ├─ Step 1: LoginDetectorModule.run(target_url)
  │    → 外部工具發現（Katana + Dirsearch）
  │    → URL 去重與預排序
  │    → LLM 並發驗證（ThreadPoolExecutor, max_workers=5）
  │    → 返回 candidates[] 和 login_result
  │
  ├─ For each candidate (並發執行):
  │    │
  │    └─ run_attack_chain(candidate_url, global_bypass_event)
  │         │
  │         ├─ Step 2: SQLiDetectorModule.run(candidate_url)
  │         │    → 提取表單參數
  │         │    → 獲取失敗基線響應
  │         │    → LLM 生成候選 Payload
  │         │    → 逐一注入 Payload 並對比
  │         │    → 若繞過成功：
  │         │         - landing_page_url = 重定向目標
  │         │         - is_authenticated = True
  │         │         - 設置 global_bypass_event（停止其他線程）
  │         │         - 固化 Session Cookie
  │         │
  │         ├─ upload_scan_url = landing_page_url 或 analysis_url
  │         │    （智能修正：若為 API 端點則退回根目錄）
  │         │
  │         ├─ Step 3: UploadIdentifierModule.run(upload_scan_url)
  │         │    → Phase 0: 首頁靜態 + 動態掃描
  │         │    → Phase 1: 關鍵詞鏈接優先探索
  │         │    → Phase 2: BFS 深度遍歷（優先編輯頁面）
  │         │    → Phase 3: Katana 動態爬蟲（帶 Session cookies）
  │         │    → Phase 4: Playwright 模擬點擊（已認證時啟用）
  │         │    → AJAX 表單推斷（LLM 分析 JS 源碼）
  │         │    → 返回 upload_id_result.findings[]
  │         │
  │         └─ Step 4: UnifiedUploadAuditModule.run(upload_scan_url)
  │              → 自動擴展：探測 action=update/edit 等變體
  │              → 對每個上傳端點：
  │                   - LLM 生成繞過策略
  │                   - 上傳無害測試文件摸底
  │                   - 上傳 Webshell 探針
  │                   - DOM 差異對比 + 網絡攔截定位路徑
  │                   - LLM 診斷與自糾錯循環
  │              → 返回 audit_result.findings[]
  │
  └─ Reporter.generate()
       → 生成 JSON/文本報告"""

    p = doc.add_paragraph()
    p.add_run(flow).font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # 5. 配置系統
    doc.add_heading('5. 配置系統', 1)

    doc.add_paragraph('所有配置集中在 config/settings.py 中，支持環境變量覆蓋：')

    config_sections = [
        ('LLM 配置', [
            ('LLM_PROVIDER', 'LLM 提供商', '"google" | "openai" | "ollama"'),
            ('LLM_MODEL', '模型名稱', '如 "gemma4-31b-gpu", "gemini-pro"'),
            ('LLM_TEMPERATURE', '溫度參數', '0.0（確定性輸出）'),
            ('OLLAMA_BASE_URL', 'Ollama 服務地址', '如 "http://192.168.1.52:1234"'),
        ]),
        ('HTTP 請求配置', [
            ('REQUEST_TIMEOUT', '請求超時秒數', '60'),
            ('REQUEST_VERIFY_SSL', 'SSL 驗證', 'False（滲透測試環境常關閉）'),
            ('REQUEST_HEADERS', '自定義請求頭', 'User-Agent 等'),
        ]),
        ('爬蟲配置', [
            ('CRAWLER_MAX_DEPTH', '登錄頁遞歸查找最大深度', '3'),
            ('UPLOAD_MAX_DEPTH', '後台尋找上傳點最大深度', '4'),
            ('UPLOAD_MAX_PAGES', '最多安全遍歷的後台頁面數量', '100'),
        ]),
        ('外部工具配置', [
            ('TOOL_DISCOVERY_ENABLED', '是否啟用外部工具', 'True'),
            ('KATANA_ENABLED', 'Katana 爬蟲開關', 'True'),
            ('KATANA_DEPTH', 'Katana 爬取深度', '3'),
            ('KATANA_TIMEOUT', 'Katana 進程總時限（秒）', '120'),
            ('DIRSEARCH_ENABLED', 'Dirsearch 開關', 'True'),
            ('DIRSEARCH_WORDLIST', '自定義字典路徑', 'login_wordlist.txt'),
        ]),
        ('Playwright 模擬點擊配置', [
            ('PLAYWRIGHT_CRAWLER_ENABLED', '是否啟用模擬點擊爬蟲', 'True'),
            ('PLAYWRIGHT_CRAWLER_MAX_PAGES', '最多探索頁面數', '50'),
            ('PLAYWRIGHT_CRAWLER_MAX_DEPTH', '最大點擊深度', '3'),
            ('PLAYWRIGHT_CRAWLER_TIMEOUT', '單頁等待超時（毫秒）', '10000'),
        ]),
    ]

    for section_name, items in config_sections:
        doc.add_heading(section_name, 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '配置項'
        hdr_cells[1].text = '說明'
        hdr_cells[2].text = '默認值/示例'

        for item_name, desc, default in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item_name
            row_cells[1].text = desc
            row_cells[2].text = default

    doc.add_page_break()

    # 6. 核心模塊詳解
    doc.add_heading('6. 核心模塊詳解', 1)

    # 6.1 Requester
    doc.add_heading('6.1 Requester - 統一 HTTP 請求封裝', 2)
    doc.add_paragraph(
        '所有模塊的 HTTP 請求都通過此類發出，集中處理 TLS 警告壓制、會話復用、'
        '統一 User-Agent、超時處理、Playwright 渲染等。'
    )

    doc.add_heading('核心方法', 3)
    methods = [
        ('get(url, **kwargs)', '發送 GET 請求，返回 Response 或 None'),
        ('post(url, data, **kwargs)', '發送 POST 請求，返回 Response 或 None'),
        ('fetch_rendered_html(url)', '使用 Playwright 獲取渲染後的 HTML（處理 SPA）'),
        ('fetch_network_resources(url)', '使用 Playwright 攔截頁面所有網絡請求'),
    ]
    for method, desc in methods:
        p = doc.add_paragraph()
        p.add_run(method).bold = True
        p.add_run(f': {desc}')

    doc.add_heading('Playwright Cookie 同步', 3)
    doc.add_paragraph(
        'Playwright 方法會自動同步 requests.Session 的 cookies 和 headers，'
        '確保在 SQL 注入繞過成功後，Playwright 也能保持已認證狀態。'
    )

    # 6.2 PageParser
    doc.add_heading('6.2 PageParser - HTML 結構化解析器', 2)
    doc.add_paragraph('負責將原始 HTML 解析為各模塊需要的結構化特徵數據。')

    doc.add_heading('核心方法', 3)
    parser_methods = [
        ('get_forms()', '提取所有表單及其輸入字段（含虛擬 AJAX 表單）'),
        ('get_upload_forms()', '篩選出包含文件上傳控件的表單'),
        ('get_auth_signals()', '提取登錄頁高信號特徵（密碼框、按鈕文字等）'),
        ('get_all_links()', '提取所有同域或相關鏈接'),
        ('get_login_candidate_links()', '提取可能導向登錄頁的候選鏈接'),
        ('get_javascript_sources()', '提取內聯和外聯 JavaScript'),
    ]
    for method, desc in parser_methods:
        p = doc.add_paragraph()
        p.add_run(method).bold = True
        p.add_run(f': {desc}')

    # 6.3 LLM Factory
    doc.add_heading('6.3 LLM Factory - LLM 工廠', 2)
    doc.add_paragraph('統一管理 LangChain LLM 實例的創建，支持多 Provider 切換。')

    doc.add_heading('核心函數', 3)
    llm_funcs = [
        ('get_llm()', '根據配置返回對應的 LangChain Chat LLM 實例'),
        ('get_structured_llm(schema)', '返回已綁定結構化輸出 schema 的 LLM'),
    ]
    for func, desc in llm_funcs:
        p = doc.add_paragraph()
        p.add_run(func).bold = True
        p.add_run(f': {desc}')

    doc.add_paragraph(
        'get_structured_llm() 會根據 Provider 自動選擇最佳方法：\n'
        '- Ollama: json_mode（本地模型依賴 format="json"）\n'
        '- Google/OpenAI: function_calling（原生支持）\n'
        '- 其他: function_calling（避免 Claude 的 assistant prefill 錯誤）'
    )

    # 6.4 Tool Discovery
    doc.add_heading('6.4 Tool Discovery - 外部工具 URL 發現層', 2)
    doc.add_paragraph(
        '整合 Katana（主動爬蟲）與 Dirsearch（目錄爆破）兩款工具，'
        '作為登錄頁識別的第一層廣度發現。'
    )

    doc.add_heading('KatanaRunner', 3)
    doc.add_paragraph(
        '封裝 Katana 主動爬蟲，優勢：\n'
        '- 支持 headless 渲染（-jc），可爬取 Vue/React SPA 動態頁面\n'
        '- 自動跟蹤 JS 注入的鏈接\n'
        '- 支持深度控制（-d）\n'
        '- 支持傳入 cookies 以已登錄狀態爬取'
    )

    doc.add_heading('DirsearchRunner', 3)
    doc.add_paragraph(
        '封裝 Dirsearch 目錄爆破工具，優勢：\n'
        '- 基於字典枚舉常見路徑（/login, /admin 等）\n'
        '- 無需頁面內有超鏈接，適合發現孤立的後台入口\n'
        '- 支持多擴展名同時爆破（php/asp/jsp 等）'
    )

    # 6.5 Playwright Interceptor
    doc.add_heading('6.5 Playwright Interceptor - 動態攔截器', 2)
    doc.add_paragraph(
        '用於處理 AJAX 表單（action 為空）的場景，通過真實點擊提交按鈕，'
        '監聽底層網絡請求，捕獲真實的後端 API 提交路徑。'
    )

    doc.add_heading('工作流程', 3)
    interceptor_flow = [
        '啟動無頭瀏覽器，注入 Session cookies',
        '加載頁面，註冊網絡請求監聽器',
        '自動填充表單（testuser/Test@1234）',
        '尋找並點擊提交按鈕（使用 JS 原生 click）',
        '攔截 POST/PUT 請求，返回真實提交地址',
    ]
    for i, step in enumerate(interceptor_flow, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    # 6.6 Playwright Crawler
    doc.add_heading('6.6 Playwright Crawler - 模擬點擊爬蟲', 2)
    doc.add_paragraph(
        'Phase 4 新增的模塊，使用真實瀏覽器模擬用戶點擊行為，'
        '遍歷後台所有可達頁面，發現上傳點。專為 SPA 框架設計。'
    )

    doc.add_heading('核心能力', 3)
    crawler_features = [
        '自動同步 requests.Session 的 cookies，保持已認證狀態',
        '智能提取可點擊元素（菜單、Tab、按鈕、鏈接）',
        '安全過濾：避免點擊刪除/登出等危險操作',
        'URL 去重：基於結構化簽名避免重複訪問',
        '上傳表單檢測：標準 input + Dropzone/WebUploader 等 JS 組件',
    ]
    for feature in crawler_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 7. 審計模塊詳解
    doc.add_heading('7. 審計模塊詳解', 1)

    # 7.1 Login Detector
    doc.add_heading('7.1 LoginDetectorModule - 登錄頁識別', 2)
    doc.add_paragraph('三層架構：工具發現 → LLM 過濾 → 遞歸補充')

    doc.add_heading('Layer 1: 廣度發現', 3)
    doc.add_paragraph(
        '使用 Katana 爬蟲和 Dirsearch 爆破發現所有可能的 URL，'
        '合併去重後進行關鍵詞預排序（login/admin/signin 優先）。'
    )

    doc.add_heading('Layer 2: LLM 精準過濾', 3)
    doc.add_paragraph(
        '對每個候選 URL 拉取頁面內容，提取特徵後交由 LLM 判斷是否為登錄頁。'
        '使用 Pydantic 結構化輸出，要求 LLM 返回：\n'
        '- is_login_page: 是否為登錄頁\n'
        '- confidence: 置信度（0.0-1.0）\n'
        '- reason: 推理過程\n'
        '- potential_login_links: 若非登錄頁，推薦的候選鏈接'
    )

    doc.add_heading('Layer 3: 遞歸補充', 3)
    doc.add_paragraph(
        '若外部工具未能發現登錄頁，或 LLM 過濾後無命中，'
        '啟動遞歸 LLM 爬蟲，從起始頁開始逐層探索。'
    )

    # 7.2 SQLi Detector
    doc.add_heading('7.2 SQLiDetectorModule - Auth Bypass 檢測', 2)
    doc.add_paragraph('基於 LLM 語義推理的身份認證繞過測試。')

    doc.add_heading('工作流程', 3)
    sqli_flow = [
        '提取頁面表單，識別登錄表單（含 password 字段）',
        '獲取基線：提交絕對錯誤的憑證，獲取「正常登錄失敗」的響應',
        'LLM 生成 Payload：根據表單結構動態生成候選 Payload',
        '注入測試：逐一注入 Payload，使用 allow_redirects=False 捕獲原始 302',
        '跟隨重定向：獲取著陸頁真實內容',
        'LLM 語義對比：對比基線與 Payload 響應的 Location、Status、內容',
        '判斷繞過：若出現「Welcome」「Admin」等已登錄特徵，判定繞過成功',
    ]
    for i, step in enumerate(sqli_flow, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_heading('關鍵設計', 3)
    doc.add_paragraph(
        '- 使用 allow_redirects=False 捕獲原始 302 響應，'
        '直接暴露 Location Header 差異\n'
        '- 跟隨重定向獲取著陸頁，讓 LLM 同時看到「302 Header」和「著陸頁文本」\n'
        '- 提取 JS Alert 彈窗（alert/layer.msg/Swal.fire 等），'
        '對老舊 PHP 系統至關重要\n'
        '- 會話固化：繞過成功後重新發送 Payload 固化 Session Cookie'
    )

    # 7.3 Upload Auditor
    doc.add_heading('7.3 UploadIdentifierModule - 上傳功能識別', 2)
    doc.add_paragraph('四階段探索 + Playwright 模擬點擊，發現所有上傳表單。')

    doc.add_heading('Phase 0: 首頁掃描', 3)
    doc.add_paragraph(
        '檢查首頁是否存在上傳表單，若靜態 HTML 未發現，'
        '使用 Playwright 動態渲染捕獲 SPA/Vue/Dropzone 等組件。'
    )

    doc.add_heading('Phase 1: 關鍵詞鏈接優先探索', 3)
    doc.add_paragraph(
        '提取頁面中包含「upload/file/avatar/profile」等關鍵詞的鏈接，'
        '優先探索這些可能包含上傳功能的頁面。'
    )

    doc.add_heading('Phase 2: BFS 深度遍歷', 3)
    doc.add_paragraph(
        '基於 UPLOAD_MAX_DEPTH 的廣度優先搜索，遍歷所有可達頁面。\n'
        '智能去重：使用 URL 簽名（保留 action 參數，歸一化 ID）避免重複。\n'
        '優先探索：編輯/更新頁面優先於普通頁面入隊。'
    )

    doc.add_heading('Phase 3: Katana 動態爬蟲', 3)
    doc.add_paragraph(
        '傳入當前 Session 的 cookies，使用 Katana 進行動態爬取，'
        '發現無超鏈接的隱藏端點。'
    )

    doc.add_heading('Phase 4: Playwright 模擬點擊', 3)
    doc.add_paragraph(
        '僅在已認證狀態下啟用。使用真實瀏覽器模擬用戶點擊菜單、Tab、按鈕，'
        '遍歷後台所有可達頁面，發現 JS 動態生成的上傳表單。'
    )

    doc.add_heading('AJAX 表單推斷', 3)
    doc.add_paragraph(
        '對於 action 為空的表單，收集頁面 JavaScript 源碼，'
        '使用 LLM 分析推斷真實的後端 API URL 和請求方法。'
    )

    # 7.4 Unified Upload Auditor
    doc.add_heading('7.4 UnifiedUploadAuditModule - Webshell 上傳驗證', 2)
    doc.add_paragraph('LLM Agent 模式：智能決策 + 自動糾錯循環。')

    doc.add_heading('自動擴展', 3)
    doc.add_paragraph(
        '從已發現的上傳點推導出同類「增/改」操作的上傳端點。\n'
        '支持 PHP/JSP/ASP/ASPX 等多種後端技術。\n'
        '探測 action=add/update/edit/avatar 等變體。'
    )

    doc.add_heading('Agent 決策鏈', 3)
    agent_flow = [
        'Strategy Agent: LLM 根據目標環境動態生成繞過策略',
        'Request Executor: Python 構建精準的 multipart/form-data 請求',
        'Path Analysis Agent: LLM 結合 DOM 差異對比解析 Webshell 真實路徑',
        'Diagnostic Agent: LLM 對 Webshell 響應進行漏洞診斷',
        'Self-Correction Loop: 若未解析或報錯，自動調整策略並重試',
    ]
    for i, step in enumerate(agent_flow, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_heading('繞過策略', 3)
    strategies = [
        'Direct Upload: 基礎直接上傳（.php）',
        'MIME Type Spoofing: Content-Type 偽造為 image/jpeg',
        'Alternative Extensions: 備選可執行後綴（.php3, .phtml, .phar）',
        'Case Obfuscation: 大小寫混合（.PhP, .pHtml）',
        'Double Extensions: 雙後綴（.jpg.php, .png.php）',
        'Null Byte / Path Traversal: 截斷或路徑穿越',
    ]
    for strategy in strategies:
        doc.add_paragraph(strategy, style='List Bullet')

    doc.add_heading('路徑推斷', 3)
    doc.add_paragraph(
        '多種策略組合：\n'
        '1. 上傳響應 JSON 解析\n'
        '2. DOM 差異對比（上傳前後頁面鏈接對比）\n'
        '3. UUID 前綴匹配（處理服務器重命名）\n'
        '4. 已知頁面掃描（文件管理器、產品列表等）\n'
        '5. 網絡攔截（Playwright 捕獲 AJAX 請求）\n'
        '6. 正則表達式兜底'
    )

    doc.add_page_break()

    # 8. 報告系統
    doc.add_heading('8. 報告系統', 1)

    doc.add_paragraph(
        'Reporter 類負責將各模塊的審計結果匯總，輸出為 JSON 或格式化文本報告。'
    )

    doc.add_heading('報告結構', 2)
    report_structure = """{
  "target": "目標 URL",
  "generated_at": "生成時間",
  "total_execution_time_seconds": 123.45,
  "global_results": [
    {
      "module": "login_detector",
      "findings": [...],
      "summary": "..."
    }
  ],
  "scanned_urls": {
    "http://target.com": [
      {
        "module": "sqli_detector",
        "findings": [...],
        "summary": "...",
        "execution_time_seconds": 10.5
      },
      {
        "module": "upload_identifier",
        "findings": [...],
        "summary": "..."
      },
      {
        "module": "unified_upload_audit",
        "findings": [...],
        "summary": "..."
      }
    ]
  }
}"""

    p = doc.add_paragraph()
    p.add_run(report_structure).font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.2)

    doc.add_heading('上傳接口清單', 2)
    doc.add_paragraph(
        '在上傳功能識別完成後，會在終端打印集中的上傳接口清單：'
    )

    list_example = """────────────────────────────────────────────────────────────
  📋 上傳接口清單 (共 3 個)
────────────────────────────────────────────────────────────
  [1] POST http://target.com/admin/upload.php
      文件參數: file, avatar
      發現頁面: http://target.com/admin/profile.php
      enctype: multipart/form-data
      允許類型: image/jpeg, image/png

  [2] POST http://target.com/admin/import.php
      文件參數: import_file
      發現頁面: http://target.com/admin/data.php
────────────────────────────────────────────────────────────"""

    p = doc.add_paragraph()
    p.add_run(list_example).font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # 9. 關鍵技術實現
    doc.add_heading('9. 關鍵技術實現', 1)

    doc.add_heading('9.1 線程安全日誌', 2)
    doc.add_paragraph(
        '使用 ThreadSafeLogger 劫持 sys.stdout，通過 threading.local '
        '為每個線程單獨隔離 stdout 內容，解決多線程日誌交織問題。'
    )

    doc.add_heading('9.2 URL 去重簽名', 2)
    doc.add_paragraph(
        '生成 URL 的結構化簽名，保留路由/動作參數（act, action, mod 等），'
        '僅對數值 ID、分頁參數（id, page, offset）進行歸一化。\n\n'
        '示例：\n'
        '- /user.php?id=123 → /user.php?id={id}\n'
        '- /page.php?act=list → /page.php?act=list（保留）\n'
        '- /page.php?act=add → /page.php?act=add（保留）'
    )

    doc.add_heading('9.3 Playwright Cookie 同步', 2)
    doc.add_paragraph(
        '將 requests.Session 的 cookies 轉換為 Playwright 格式並注入：\n\n'
        '代碼示例：\n'
        'pw_cookies = []\n'
        'for c in self.session.cookies:\n'
        '    pw_cookies.append({\n'
        '        "name": c.name,\n'
        '        "value": c.value,\n'
        '        "domain": c.domain or domain,\n'
        '        "path": c.path or "/"\n'
        '    })\n'
        'context.add_cookies(pw_cookies)'
    )

    doc.add_heading('9.4 LLM 結構化輸出', 2)
    doc.add_paragraph(
        '使用 Pydantic 模型定義輸出結構，通過 with_structured_output() 綁定：\n\n'
        'class AuthBypassResult(BaseModel):\n'
        '    is_bypassed: bool\n'
        '    confidence: str  # "high" | "medium" | "low"\n'
        '    bypass_evidence: List[str]\n'
        '    reason: str\n\n'
        'chain = PROMPT | get_structured_llm(AuthBypassResult)\n'
        'result: AuthBypassResult = chain.invoke({...})'
    )

    doc.add_heading('9.5 智能降級', 2)
    doc.add_paragraph(
        '系統設計了多層降級機制，確保任意組件失敗不影響整體流水線：\n'
        '- Playwright 未安裝 → 降級為普通 requests\n'
        '- Katana 未安裝 → 跳過，使用其他發現方式\n'
        '- Dirsearch 未安裝 → 跳過，使用其他發現方式\n'
        '- LLM 調用失敗 → 使用正則表達式備用方案\n'
        '- 外部工具無結果 → 啟動遞歸 LLM 爬蟲'
    )

    doc.add_page_break()

    # 10. 使用說明
    doc.add_heading('10. 使用說明', 1)

    doc.add_heading('10.1 安裝依賴', 2)
    install_code = """pip install -r requirements.txt
playwright install chromium"""

    p = doc.add_paragraph()
    p.add_run(install_code).font.name = 'Courier New'

    doc.add_heading('10.2 基本用法', 2)

    doc.add_paragraph('單目標掃描：')
    cmd1 = 'python -m web_audit.main --url https://target.example.com'
    p = doc.add_paragraph()
    p.add_run(cmd1).font.name = 'Courier New'

    doc.add_paragraph('多目標批量掃描：')
    cmd2 = 'python -m web_audit.main -f targets.txt -t 5'
    p = doc.add_paragraph()
    p.add_run(cmd2).font.name = 'Courier New'

    doc.add_paragraph('單獨運行某個模塊：')
    cmd3 = 'python -m web_audit.main --url https://target.com --step login'
    p = doc.add_paragraph()
    p.add_run(cmd3).font.name = 'Courier New'

    doc.add_heading('10.3 可用步驟', 2)
    steps = [
        ('all', '運行完整流水線（默認）'),
        ('login', '僅運行登錄頁識別'),
        ('sqli', '僅運行 SQL 注入繞過檢測'),
        ('upload_id', '僅運行上傳功能識別'),
        ('upload_audit', '僅運行上傳安全審查'),
        ('upload_exploit', '僅運行 Webshell 漏洞驗證'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '步驟'
    hdr_cells[1].text = '說明'
    for step, desc in steps:
        row_cells = table.add_row().cells
        row_cells[0].text = step
        row_cells[1].text = desc

    doc.add_heading('10.4 環境變量', 2)
    env_vars = [
        ('AUDIT_TARGET_URL', '目標 URL（可替代 --url 參數）'),
        ('AUDIT_COOKIE', '自定義 Cookie（注入到 Session）'),
        ('GOOGLE_API_KEY', 'Google Gemini API Key'),
        ('OPENAI_API_KEY', 'OpenAI API Key'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '環境變量'
    hdr_cells[1].text = '說明'
    for var, desc in env_vars:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = desc

    doc.add_heading('10.5 報告輸出', 2)
    doc.add_paragraph(
        '報告默認輸出到 ./reports/ 目錄，格式為 JSON。\n'
        '文件名格式：audit_{target}_{timestamp}.json\n\n'
        '可在 config/settings.py 中修改：\n'
        '- REPORT_OUTPUT_DIR: 報告輸出目錄\n'
        '- REPORT_FORMAT: 報告格式（"json" | "text"）'
    )

    # 保存文檔
    output_path = os.path.join(os.path.dirname(__file__), 'Web_Audit_架構文檔.docx')
    doc.save(output_path)
    print(f'✅ 文檔已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
